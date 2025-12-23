"""
Development Services Proposal Generator
Streamlit web application for generating professional proposal documents
With Kimley-Horn header and footer
"""

import streamlit as st
from datetime import date
from io import BytesIO
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ROW_HEIGHT_RULE, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

# ============================================================================
# TASK DESCRIPTIONS DATABASE
# ============================================================================

DEFAULT_FEES = {
    '110': {'name': 'Civil Engineering Design', 'amount': 40000, 'type': 'Hourly, Not-to-Exceed'},
    '120': {'name': 'Civil Schematic Design', 'amount': 35000, 'type': 'Hourly, Not-to-Exceed'},
    '130': {'name': 'Civil Design Development', 'amount': 45000, 'type': 'Hourly, Not-to-Exceed'},
    '140': {'name': 'Civil Construction Documents', 'amount': 50000, 'type': 'Hourly, Not-to-Exceed'},
    '150': {'name': 'Civil Permitting', 'amount': 40000, 'type': 'Hourly, Not-to-Exceed'},
    '210': {'name': 'Meetings and Coordination', 'amount': 20000, 'type': 'Hourly, Not-to-Exceed'},
    '310': {'name': 'Civil Construction Phase Services', 'amount': 35000, 'type': 'Lump Sum'}
}

TASK_DESCRIPTIONS = {
    '110': [
        "Kimley-Horn will prepare an onsite drainage report with supporting calculations showing the proposed development plan is consistent with the Southwest Florida Water Management District Basis of Review. This design will account for the stormwater design to support the development of the project site. The drainage report will include limited stormwater modeling to demonstrate that the Lot A site development will maintain the existing discharge rate and provide the required stormwater attenuation.",
        "The onsite drainage report will include calculations for 25-year 24-hour and 100-year 24-hour design storm conditions in accordance with Southwest Florida Water Management District Guidelines. A base stormwater design will be provided for the project site showing reasonable locations for stormwater conveyance features and stormwater management pond sizing."
    ],
    '120': [
        "Kimley-Horn will prepare Civil Schematic Design deliverables in accordance with the Client's Design Project Deliverables Checklist. For the Civil Schematic Design task, the deliverables that Kimley-Horn will provide consist of Civil Site Plan, Establish Finish Floor Elevations, Utility Will Serve Letters and Points of Service, Utility Routing and Easement Requirements."
    ],
    '130': [
        "Upon Client approval of the Schematic Design task, Kimley-Horn will prepare Design Development Plans of the civil design in accordance with the Client's Design Project Deliverables Checklist for Civil Design Development Deliverables. These documents will be approximately 50% complete and will include detail for City code review and preliminary pricing but will not include enough detail for construction bidding."
    ],
    '140': [
        "Based on the approved Development Plan, Kimley-Horn will provide engineering and design services for the preparation of site construction plans for on-site improvements.",
        "Cover Sheet",
        "The cover sheet includes plan contents, vicinity map, legal description and team identification.",
        "Existing Conditions Plan/Demolition Plan",
        "This sheet will include and identify the required demolition of the existing items on the project site.",
        "Site Layout Plan",
        "This sheet will include building setback lines, property lines, outline of building footprint, parking areas, handicap access ramps, sidewalks, crosswalks, driveways, and traffic lanes.",
        "Grading and Drainage Plan",
        "This sheet will include existing and proposed spot elevations and contours, building finish floor elevations, parking area drainage patterns, and stormwater inlet and pipe locations and sizes.",
        "**NOTE:** Any structural retaining walls are not included with this scope and shall be designed and permitted by others.",
        "Utility Plan",
        "This sheet will show the location and size of all water, sanitary sewer and reclaimed water facilities required to serve the development.",
        "**NOTE:** Kimley-Horn's contract does not include the design of the fire lines from the designated point of service (P.O.S.) up to 1' above the building foundation.",
        "Erosion and Sediment Control Plan",
        "This sheet will include erosion and sediment control measures designed to be implemented during construction.",
        "Details",
        "Standard and modified typical construction details will be provided.",
        "**NOTE:** A specifications package is not included in this scope of services as specifications are per authority having jurisdiction (AHJ)."
    ],
    '150': [
        "Prepare and submit on the Client's behalf the following permitting packages for review/approval of construction documents, and attend meetings required to obtain the following Agency approvals:",
        "Southwest Florida Water Management District Environmental Resource Permit – Minor Modification",
        "City of Tampa Water Department Commitment / Construction Plan Approval",
        "Hillsborough County Environmental Protection Commission",
        "Kimley-Horn will coordinate with the City of Tampa Development Review and coordination with the Florida Department of Transportation and the Hillsborough County departments as needed to obtain the necessary regulatory and utility approval of the site plans and associated drainage facilities. We will assist the Client with meetings necessary to gain site plan approval.",
        "This scope does not anticipate a Geotechnical or Environmental Assessment Report, Survey, Topographic Survey, or Arborist Report be required for this permit application.",
        "It is assumed Client will provide the needed information regarding the development program and requirements. Kimley-Horn will work with the Owner and their team to integrate the necessary design requirements into the Civil design to support entitlement, platting, and development approvals.",
        "These permit applications will be submitted using the electronic permitting submittal system (web-based system) for the respective jurisdictions where applicable."
    ],
    '210': [
        "Kimley-Horn will be available to provide miscellaneous project support at the direction of the Client. This task may include design meetings, additional permit support, permit research, or other miscellaneous tasks associated with the initial and future development of the project site. This task will also cover tasks such as design coordination meetings, scheduling, coordination with other client consultants, responses to additional rounds of agency comments."
    ],
    '310': [
        "Engineering construction phase services will be performed in connection with site improvements designed by Kimley-Horn. The scope of this task assumes construction phase services will be performed concurrent and in coordination with one General Contractor for the entire project. This task does not include constructing the project in multiple phases. Kimley-Horn construction phase services will include the following:",
        "Provide for review of shop drawings and submittals required for the site improvements controlled by our design documents. Kimley-Horn has included up to {shop_drawing_hours} hours for review of shop drawings and samples.",
        "Review and reply to Contractor's request(s) for information during construction phase. Kimley-Horn has included up to {rfi_hours} hours for response to RFI's.",
        "Attendance at up to {oac_meetings} one-hour each Owner-Architect-Contractor (OAC) virtual meetings.",
        "Kimley-Horn will visit the construction site during the duration of construction for an estimated total of up to {site_visits} site visits at two-hours each to observe the progress of the civil components of work completed.",
        "Provide up to two (2) reviews of 'as-built' documents, submitted by General Contractor's registered land surveyor.",
        "Kimley-Horn will prepare Record Drawings for potable water and sanitary sewer only. Kimley-Horn has included up to {record_drawing_hours} hours for record drawing preparation.",
        "Kimley-Horn will submit FDEP water and sewer clearance submittals based on as-built information provided by the Contractor.",
        "Kimley-Horn shall submit a Letter of General Compliance for the civil related components of construction to the AHJ.",
        "Submit Certification of Completion to the Water Management District (WMD).",
        "The above hours allocated to the respective construction phase services may be interchangeable amongst the construction phase services outlined in this task, however the total number of hours included within the entirety of the task is up to {total_hours} hours."
    ]
}

# ============================================================================
# PERMIT CONFIGURATION BY COUNTY
# ============================================================================

PERMIT_MAPPING = {
    "Pinellas": {
        "ahj_name": "Pinellas County",
        "wmd": "Southwest Florida Water Management District",
        "wmd_short": "SWFWMD",
        "default_permits": ["ahj", "wmd_erp", "sewer", "water"]
    },
    "Hillsborough": {
        "ahj_name": "Hillsborough County",
        "wmd": "Southwest Florida Water Management District", 
        "wmd_short": "SWFWMD",
        "default_permits": ["ahj", "wmd_erp", "sewer", "water"]
    },
    "Pasco": {
        "ahj_name": "Pasco County",
        "wmd": "Southwest Florida Water Management District",
        "wmd_short": "SWFWMD", 
        "default_permits": ["ahj", "wmd_erp", "sewer", "water"]
    },
    "Manatee": {
        "ahj_name": "Manatee County",
        "wmd": "Southwest Florida Water Management District",
        "wmd_short": "SWFWMD",
        "default_permits": ["ahj", "wmd_erp", "sewer", "water"]
    },
    "Sarasota": {
        "ahj_name": "Sarasota County",
        "wmd": "Southwest Florida Water Management District",
        "wmd_short": "SWFWMD",
        "default_permits": ["ahj", "wmd_erp", "sewer", "water"]
    },
    "Polk": {
        "ahj_name": "Polk County",
        "wmd": "Southwest Florida Water Management District",
        "wmd_short": "SWFWMD",
        "default_permits": ["ahj", "wmd_erp", "sewer", "water"]
    }
}

# ============================================================================
# PROPERTY LOOKUP FUNCTIONS
# ============================================================================

@st.cache_data(ttl=3600, show_spinner=False)
def lookup_hillsborough_property(parcel_id):
    """
    Lookup property from Hillsborough County via SWFWMD regional service.
    FOLIO format: Remove dash (192605-0030 becomes 1926050030)
    """
    import requests
    
    # Use SWFWMD regional service - more reliable
    base_url = "https://www25.swfwmd.state.fl.us/arcgis12/rest/services/BaseVector/parcel_search/MapServer/7/query"
    
    # Remove dash
    folio_normalized = parcel_id.replace('-', '')
    
    params = {
        'where': f"FOLIONUM='{folio_normalized}'",
        'outFields': '*',
        'returnGeometry': 'false',
        'f': 'json'
    }
    
    try:
        response = requests.get(base_url, params=params, timeout=15)
        response.raise_for_status()
        data = response.json()
        
        if data.get('features') and len(data['features']) > 0:
            attr = data['features'][0]['attributes']
            
            return {
                'success': True,
                'address': attr.get('SITUSADD1', ''),
                'city': attr.get('SCITY', 'TAMPA'),
                'zip': attr.get('SZIP', ''),
                'owner': attr.get('OWNERNAME', ''),
                'land_use': attr.get('PARUSEDESC', ''),
                'zoning': 'Contact City/County for zoning info',
                'error': None
            }
        else:
            return {
                'success': False,
                'error': f'Parcel {folio_normalized} not found in database'
            }
    
    except Exception as e:
        return {'success': False, 'error': f'API Error: {str(e)}'}

@st.cache_data(ttl=3600, show_spinner=False)
def lookup_manatee_property(parcel_id):
    """
    Lookup property from Manatee County.
    Excellent data completeness - all fields in single layer.
    """
    import requests
    
    base_url = "https://www.mymanatee.org/gisits/rest/services/opendata/Planning/MapServer/22/query"
    
    params = {
        'where': f"PIN='{parcel_id}'",
        'outFields': '*',
        'returnGeometry': 'false',
        'f': 'json'
    }
    
    try:
        response = requests.get(base_url, params=params, timeout=10)
        response.raise_for_status()
        data = response.json()
        
        if data.get('features') and len(data['features']) > 0:
            attr = data['features'][0]['attributes']
            
            return {
                'success': True,
                'address': attr.get('PRIMARY_ADDRESS', ''),
                'city': attr.get('PROP_CITYNAME', ''),
                'zip': attr.get('PROP_ZIP', ''),
                'owner': attr.get('OWNER', ''),
                'land_use': attr.get('FUTURE_LAND_USE', ''),
                'zoning': attr.get('ZONING', ''),
                'error': None
            }
        else:
            return {'success': False, 'error': 'Parcel ID not found in Manatee County'}
    
    except Exception as e:
        return {'success': False, 'error': f'API Error: {str(e)}'}


def lookup_property_info(county, parcel_id):
    """Lookup property info based on county."""
    if county == "Hillsborough":
        return lookup_hillsborough_property(parcel_id)
    elif county == "Manatee":
        return lookup_manatee_property(parcel_id)
    else:
        return {'success': False, 'error': f'{county} County lookup not yet implemented. Please enter information manually.'}

# ============================================================================
# DOCUMENT GENERATION FUNCTIONS
# ============================================================================

def set_cell_background(cell, color_hex):
    """Set cell background color."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    existing_shd = tcPr.find(qn('w:shd'))
    if existing_shd is not None:
        tcPr.remove(existing_shd)
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), color_hex)
    tcPr.append(shd)


def set_cell_margins(cell, top=20, bottom=20, start=40, end=40):
    """Set cell margins in twips."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    existing = tcPr.find(qn('w:tcMar'))
    if existing is not None:
        tcPr.remove(existing)
    tcMar = OxmlElement('w:tcMar')
    for margin_name, value in [('top', top), ('bottom', bottom), ('start', start), ('end', end)]:
        margin = OxmlElement(f'w:{margin_name}')
        margin.set(qn('w:w'), str(value))
        margin.set(qn('w:type'), 'dxa')
        tcMar.append(margin)
    tcPr.append(tcMar)


def remove_table_borders(table):
    """Remove all borders from table."""
    tbl = table._tbl
    tblPr = tbl.tblPr
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        tbl.insert(0, tblPr)
    existing = tblPr.find(qn('w:tblBorders'))
    if existing is not None:
        tblPr.remove(existing)
    tblBorders = OxmlElement('w:tblBorders')
    for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'nil')
        tblBorders.append(border)
    tblPr.append(tblBorders)


def create_header(section):
    """Create Kimley-Horn header."""
    header = section.header
    header.is_linked_to_previous = False
    
    header_table = header.add_table(rows=1, cols=2, width=Inches(6.5))
    header_table.autofit = False
    header_table.columns[0].width = Inches(5.0)
    header_table.columns[1].width = Inches(1.5)
    
    tbl = header_table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else OxmlElement('w:tblPr')
    tblBorders = OxmlElement('w:tblBorders')
    for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'none')
        tblBorders.append(border)
    tblPr.append(tblBorders)
    if tbl.tblPr is None:
        tbl.insert(0, tblPr)
    
    logo_cell = header_table.cell(0, 0)
    logo_cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    logo_para = logo_cell.paragraphs[0]
    logo_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    logo_para.clear()
    
    run1 = logo_para.add_run("Kimley")
    run1.font.size = Pt(28)
    run1.font.bold = False
    run1.font.color.rgb = RGBColor(88, 89, 91)
    run1.font.name = 'Arial Narrow'
    
    run2 = logo_para.add_run("»")
    run2.font.size = Pt(28)
    run2.font.bold = False
    run2.font.color.rgb = RGBColor(88, 89, 91)
    run2.font.name = 'Arial Narrow'
    
    run3 = logo_para.add_run("Horn")
    run3.font.size = Pt(28)
    run3.font.bold = False
    run3.font.color.rgb = RGBColor(166, 25, 46)
    run3.font.name = 'Arial Narrow'
    
    page_cell = header_table.cell(0, 1)
    page_cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    page_para = page_cell.paragraphs[0]
    page_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    
    run = page_para.add_run('Page ')
    run.font.name = 'Arial'
    run.font.size = Pt(11)
    run.font.italic = True
    run.font.color.rgb = RGBColor(0, 0, 0)
    
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = 'PAGE'
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'end')
    run._r.append(fldChar1)
    run._r.append(instrText)
    run._r.append(fldChar2)


def create_footer(section):
    """Create Kimley-Horn footer."""
    footer = section.footer
    footer.is_linked_to_previous = False
    
    widths = [Inches(1.1), Inches(0.01), Inches(4.23), Inches(0.01), Inches(0.96)]
    colors = ['5F5F5F', None, 'A20C33', None, 'A20C33']
    texts = ['kimley-horn.com', '', '200 Central Avenue Suite 600 St. Petersburg, FL 33701', '', '(727) 822-5150']
    
    table = footer.add_table(rows=1, cols=5, width=sum(widths))
    table.allow_autofit = False
    remove_table_borders(table)
    
    row = table.rows[0]
    row.height = Inches(0.22)
    row.height_rule = WD_ROW_HEIGHT_RULE.EXACTLY
    
    for idx, cell in enumerate(row.cells):
        table.columns[idx].width = widths[idx]
        cell.width = widths[idx]
        
        if colors[idx]:
            set_cell_background(cell, colors[idx])
        
        set_cell_margins(cell, top=20, bottom=20, start=40, end=40)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        
        if texts[idx]:
            para = cell.paragraphs[0]
            para.paragraph_format.space_before = Pt(0)
            para.paragraph_format.space_after = Pt(0)
            para.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
            para.clear()
            
            run = para.add_run(texts[idx])
            run.font.name = 'Arial'
            run.font.size = Pt(8)
            run.font.color.rgb = RGBColor(255, 255, 255)


def add_opening_section(doc, client_info, project_info):
    """Add opening section."""
    
    date_para = doc.add_paragraph()
    date_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = date_para.add_run(project_info['date'])
    run.font.name = 'Arial'
    run.font.size = Pt(10)
    date_para.paragraph_format.space_after = Pt(0)
    date_para.paragraph_format.line_spacing = 1.0
    
    doc.add_paragraph()
    
    para = doc.add_paragraph()
    run = para.add_run(client_info['contact'])
    run.font.name = 'Arial'
    run.font.size = Pt(10)
    para.paragraph_format.space_after = Pt(0)
    para.paragraph_format.line_spacing = 1.0
    
    para = doc.add_paragraph()
    run = para.add_run(client_info['name'])
    run.font.name = 'Arial'
    run.font.size = Pt(10)
    para.paragraph_format.space_after = Pt(0)
    para.paragraph_format.line_spacing = 1.0
    
    para = doc.add_paragraph()
    run = para.add_run(client_info['address1'])
    run.font.name = 'Arial'
    run.font.size = Pt(10)
    para.paragraph_format.space_after = Pt(0)
    para.paragraph_format.line_spacing = 1.0
    
    para = doc.add_paragraph()
    run = para.add_run(client_info['address2'])
    run.font.name = 'Arial'
    run.font.size = Pt(10)
    para.paragraph_format.space_after = Pt(0)
    para.paragraph_format.line_spacing = 1.0
    
    doc.add_paragraph()
    
    para = doc.add_paragraph()
    run = para.add_run('Re:\tProfessional Services Agreement')
    run.font.name = 'Arial'
    run.font.size = Pt(10)
    para.paragraph_format.space_after = Pt(0)
    para.paragraph_format.line_spacing = 1.0
    
    para = doc.add_paragraph()
    run = para.add_run(f'\t{project_info["name"]}')
    run.font.name = 'Arial'
    run.font.size = Pt(10)
    para.paragraph_format.space_after = Pt(0)
    para.paragraph_format.line_spacing = 1.0
    
    doc.add_paragraph()
    
    para = doc.add_paragraph()
    run = para.add_run(f'Dear {client_info["contact"].split()[0]} {client_info["contact"].split()[-1]}:')
    run.font.name = 'Arial'
    run.font.size = Pt(10)
    run.font.bold = False  # NOT bold
    para.paragraph_format.space_after = Pt(0)
    para.paragraph_format.line_spacing = 1.0
    
    doc.add_paragraph()
    
    para = doc.add_paragraph()
    opening_text = f'Kimley-Horn and Associates, Inc. ("Kimley-Horn" or "Consultant") is pleased to submit this Professional Services Agreement ("Agreement") to {client_info["name"]} ("Client") for professional services for the {project_info["name"]} ("Project").'
    run = para.add_run(opening_text)
    run.font.name = 'Arial'
    run.font.size = Pt(10)
    run.font.bold = True
    para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    para.paragraph_format.space_after = Pt(0)
    para.paragraph_format.line_spacing = 1.0
    
    doc.add_paragraph()


def add_project_understanding(doc, project_description, assumptions):
    """Add Project Understanding section with assumptions."""
    
    # Section heading - BOLD + CENTERED
    para = doc.add_paragraph()
    run = para.add_run('PROJECT UNDERSTANDING')
    run.font.name = 'Arial'
    run.font.size = Pt(11)
    run.font.bold = True  # BOLD
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    para.paragraph_format.space_after = Pt(0)
    para.paragraph_format.line_spacing = 1.0
    
    doc.add_paragraph()
    
    # Project description - JUSTIFIED
    para = doc.add_paragraph()
    run = para.add_run(project_description)
    run.font.name = 'Arial'
    run.font.size = Pt(10)
    para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    para.paragraph_format.space_after = Pt(0)
    para.paragraph_format.line_spacing = 1.0
    
    doc.add_paragraph()
    
    # Assumptions intro
    para = doc.add_paragraph()
    run = para.add_run('Kimley-Horn understands the following in preparing this proposal:')
    run.font.name = 'Arial'
    run.font.size = Pt(10)
    para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    para.paragraph_format.space_after = Pt(0)
    para.paragraph_format.line_spacing = 1.0
    
    doc.add_paragraph()
    
    # Assumptions as bullet points
    for assumption in assumptions:
        para = doc.add_paragraph(style='List Bullet')
        run = para.add_run(assumption)
        run.font.name = 'Arial'
        run.font.size = Pt(10)
        para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        para.paragraph_format.space_after = Pt(0)
        para.paragraph_format.line_spacing = 1.0
    
    doc.add_paragraph()
    
    # Closing statement
    para = doc.add_paragraph()
    run = para.add_run('If any of these assumptions are not correct, then the scope and fee will change. Based on the above understanding, Kimley-Horn proposes the following scope of services:')
    run.font.name = 'Arial'
    run.font.size = Pt(10)
    para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    para.paragraph_format.space_after = Pt(0)
    para.paragraph_format.line_spacing = 1.0
    
    doc.add_paragraph()


def add_scope_of_services(doc, selected_tasks, permits):
    """Add Scope of Services section."""
    
    # Section heading - BOLD + CENTERED
    para = doc.add_paragraph()
    run = para.add_run('Scope of Services')
    run.font.name = 'Arial'
    run.font.size = Pt(11)
    run.font.bold = True  # BOLD
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    para.paragraph_format.space_after = Pt(0)
    para.paragraph_format.line_spacing = 1.0
    
    doc.add_paragraph()
    
    para = doc.add_paragraph()
    run = para.add_run('Kimley-Horn will provide the services specifically set forth below.')
    run.font.name = 'Arial'
    run.font.size = Pt(10)
    para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    para.paragraph_format.space_after = Pt(0)
    para.paragraph_format.line_spacing = 1.0
    
    doc.add_paragraph()
    
    sub_section_keywords = ['cover sheet', 'utility plan', 'site layout', 'site plan',
                           'grading plan', 'drainage plan', 'paving', 'erosion control',
                           'detail', 'existing conditions', 'demolition']
    
    for task_num in sorted(selected_tasks.keys()):
        task = selected_tasks[task_num]
        
        # Special handling for Task 310 - Construction Phase Services
        if task_num == '310' and 'hours' in task:
            hours = task['hours']
            descriptions = []
            for desc in TASK_DESCRIPTIONS[task_num]:
                # Replace hour placeholders
                desc = desc.replace('{shop_drawing_hours}', str(hours['shop_drawing']))
                desc = desc.replace('{rfi_hours}', str(hours['rfi']))
                desc = desc.replace('{oac_meetings}', str(hours['oac_meetings']))
                desc = desc.replace('{site_visits}', str(hours['site_visits']))
                desc = desc.replace('{record_drawing_hours}', str(hours['record_drawing']))
                desc = desc.replace('{total_hours}', str(hours['total']))
                descriptions.append(desc)
        # Special handling for Task 150 - Civil Permitting
        elif task_num == '150' and permits:
            descriptions = [
                "Kimley-Horn will prepare and submit on the Client's behalf the civil construction documents to the following agencies:"
            ]
            # Add permits as separate items for bullet point formatting
            descriptions.extend(permits)
            descriptions.extend([
                "Includes development of a Stormwater Ownership and Maintenance (O&M) Plan",
                "Kimley-Horn will respond to a maximum of three (3) requests for information during the agency review process for obtaining the above permits.",
                "**BOLD:**Permit fees and impact fees are not included. Kimley-Horn does not guarantee the issuance of permits or approvals."
            ])
        else:
            descriptions = TASK_DESCRIPTIONS[task_num]
        
        # Task heading - BOLD + UNDERLINED (no automatic page break)
        para = doc.add_paragraph()
        # Removed: para.paragraph_format.page_break_before = True
        run = para.add_run(f'Task {task_num} – {task["name"].replace("Civil ", "")}')
        run.font.name = 'Arial'
        run.font.size = Pt(10)
        run.font.bold = True
        run.font.underline = True
        para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        para.paragraph_format.space_after = Pt(0)
        para.paragraph_format.line_spacing = 1.0
        
        doc.add_paragraph()
        
        permit_list_started = False
        
        for desc in descriptions:
            # Check if this is a permit item (short, from permits list)
            is_permit_bullet = (task_num == '150' and desc in permits)
            
            # Start bullet list after intro
            if task_num == '150' and 'following agencies:' in desc:
                para = doc.add_paragraph()
                run = para.add_run(desc)
                run.font.name = 'Arial'
                run.font.size = Pt(10)
                para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                para.paragraph_format.space_after = Pt(0)
                para.paragraph_format.line_spacing = 1.0
                doc.add_paragraph()
                permit_list_started = True
                continue
            
            if is_permit_bullet:
                # Permit as bullet point - INDENTED
                para = doc.add_paragraph(style='List Bullet')
                para.paragraph_format.left_indent = Inches(0.25)  # Indent for clarity
                run = para.add_run(desc)
                run.font.name = 'Arial'
                run.font.size = Pt(10)
                para.alignment = WD_ALIGN_PARAGRAPH.LEFT
                para.paragraph_format.space_after = Pt(0)
                para.paragraph_format.line_spacing = 1.0
                continue
            
            # End permit bullets, back to regular paragraphs
            if permit_list_started and not is_permit_bullet:
                permit_list_started = False
                doc.add_paragraph()
            
            # Check if bold text
            is_bold_para = desc.startswith('**BOLD:**')
            
            para = doc.add_paragraph()
            
            # Check if sub-section heading
            is_subsection = (len(desc) < 100 and 
                           any(kw in desc.lower() for kw in sub_section_keywords) and
                           not desc.endswith('.') and
                           not desc.startswith('**NOTE:**') and
                           not desc.startswith('**BOLD:**'))
            
            # Check if it's a note
            is_note = desc.startswith('**NOTE:**')
            
            if is_note:
                run = para.add_run(desc.replace('**NOTE:**', 'Note:'))
                run.font.name = 'Arial'
                run.font.size = Pt(10)
                run.font.bold = True
                run.font.italic = True
            elif is_bold_para:
                run = para.add_run(desc.replace('**BOLD:**', ''))
                run.font.name = 'Arial'
                run.font.size = Pt(10)
                run.font.bold = True
            else:
                run = para.add_run(desc)
                run.font.name = 'Arial'
                run.font.size = Pt(10)
                
                if is_subsection:
                    run.font.underline = True
            
            para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            para.paragraph_format.space_after = Pt(0)
            para.paragraph_format.line_spacing = 1.0
            
            if not is_subsection:
                doc.add_paragraph()


def add_scope_table(doc, selected_tasks):
    """Add Scope of Work table."""
    
    total_fee = sum(task['fee'] for task in selected_tasks.values())
    
    num_rows = len(selected_tasks) + 2
    table = doc.add_table(rows=num_rows, cols=4)
    table.style = 'Light Grid Accent 1'
    
    header_cells = table.rows[0].cells
    header_cells[0].text = 'Task Number & Name'
    header_cells[1].text = 'Task Number & Name'
    header_cells[2].text = 'Fee'
    header_cells[3].text = 'Type'
    
    for cell in header_cells:
        cell.paragraphs[0].runs[0].font.bold = True
        cell.paragraphs[0].runs[0].font.size = Pt(10)
        cell.paragraphs[0].runs[0].font.name = 'Arial'
    
    for idx, (task_num, task) in enumerate(sorted(selected_tasks.items()), start=1):
        row = table.rows[idx]
        row.cells[0].text = task_num
        row.cells[1].text = task['name']
        row.cells[2].text = f'$ {task["fee"]:,}'
        row.cells[3].text = task['type']
        
        for cell in row.cells:
            cell.paragraphs[0].runs[0].font.size = Pt(10)
            cell.paragraphs[0].runs[0].font.name = 'Arial'
    
    total_row = table.rows[-1]
    total_row.cells[0].text = 'Total'
    total_row.cells[1].text = 'Total'
    total_row.cells[2].text = f'$ {total_fee:,}'
    total_row.cells[3].text = f'$ {total_fee:,}'
    
    for cell in total_row.cells:
        cell.paragraphs[0].runs[0].font.bold = True
        cell.paragraphs[0].runs[0].font.size = Pt(10)
        cell.paragraphs[0].runs[0].font.name = 'Arial'
    
    doc.add_paragraph()
    
    task_list = ', '.join(sorted(selected_tasks.keys()))
    para = doc.add_paragraph()
    run = para.add_run(f'Kimley-Horn will perform the services in Tasks {task_list} on a labor fee plus expense basis with the maximum labor fee shown above.')
    run.font.name = 'Arial'
    run.font.size = Pt(10)
    para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    para.paragraph_format.space_after = Pt(0)
    para.paragraph_format.line_spacing = 1.0


def generate_proposal_document(client_info, project_info, selected_tasks, assumptions, permits, output_path):
    """Generate complete proposal document."""
    
    doc = Document()
    
    section = doc.sections[0]
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)
    
    create_header(section)
    create_footer(section)
    
    add_opening_section(doc, client_info, project_info)
    add_project_understanding(doc, project_info['description'], assumptions)
    add_scope_of_services(doc, selected_tasks, permits)
    add_scope_table(doc, selected_tasks)
    
    doc.save(output_path)
    return output_path

# ============================================================================
# STREAMLIT APP
# ============================================================================

st.set_page_config(
    page_title="Development Services Proposal Generator",
    page_icon="🏗️",
    layout="wide"
)

st.title("🏗️ Development Services Proposal Generator")
st.markdown("---")

# Create tabs
tab1, tab2, tab3, tab4, tab5 = st.tabs([
    "📍 Property & Client",
    "📋 Project Details",
    "✅ Scope of Services",
    "📧 Permitting & Billing",
    "📄 Generate Document"
])

# TAB 1: Property & Client Information
with tab1:
    # Section 1: Property/Parcel Information
    st.header("📍 Property/Parcel Information")
col_prop1, col_prop2, col_prop3 = st.columns(3)

with col_prop1:
    county = st.selectbox(
        "County *",
        options=["", "Pinellas", "Hillsborough", "Pasco", "Manatee", "Sarasota", "Polk"],
        help="Select the county where the project is located",
        key="county"
    )

with col_prop2:
    city = st.text_input(
        "City *",
        placeholder="e.g., St. Petersburg",
        key="city"
    )

with col_prop3:
    parcel_id = st.text_input(
        "Parcel ID",
        placeholder="e.g., 12-34-56-78900-000-0000",
        help="Property parcel identification number",
        key="parcel_id"
    )

# Lookup button
if st.button("🔍 Lookup Property Info", disabled=not (st.session_state.get('county') and st.session_state.get('parcel_id'))):
    with st.spinner(f"Looking up property info in {st.session_state.get('county', '')} County..."):
        result = lookup_property_info(st.session_state.get('county', ''), st.session_state.get('parcel_id', ''))
        
        if result['success']:
            st.success(f"✅ Property found!")
            st.session_state['lookup_address'] = result.get('address', '')
            st.session_state['lookup_city'] = result.get('city', '')
            st.session_state['lookup_zip'] = result.get('zip', '')
            st.session_state['lookup_owner'] = result.get('owner', '')
            st.session_state['lookup_land_use'] = result.get('land_use', '')
            st.session_state['lookup_zoning'] = result.get('zoning', '')
            
            # Display found info
            st.info(f"""
            **Found Property Information:**
            - Address: {result.get('address', 'N/A')}
            - City: {result.get('city', 'N/A')}, FL {result.get('zip', 'N/A')}
            - Owner: {result.get('owner', 'N/A')}
            - Land Use: {result.get('land_use', 'N/A')}
            - Zoning: {result.get('zoning', 'N/A')}
            
            *Fields below will be auto-filled. Please verify accuracy.*
            """)
            st.rerun()
        else:
            st.error(f"❌ {result['error']}")

    st.markdown("---")
    
    # Client Information (still in tab1)
    st.header("📋 Client Information")
    col1, col2 = st.columns(2)
    
    with col1:
        client_name = st.text_input("Client Name *", placeholder="e.g., ABC Development Corporation", key="client_name")
        address_line1 = st.text_input("Address Line 1 *", placeholder="e.g., 123 Main Street", key="address_line1")
        address_line2 = st.text_input("Address Line 2 *", placeholder="e.g., Tampa, FL 33602", key="address_line2")
    
    with col2:
        contact_person = st.text_input("Contact Person *", placeholder="e.g., Ms. Michelle Bach", key="contact_person")
        phone = st.text_input("Phone Number", placeholder="e.g., (813) 555-1234", key="phone")
        email = st.text_input("Email Address", placeholder="e.g., info@example.com", key="email")

# TAB 2: Project Details
with tab2:
    # Project Details
    st.header("📍 Project Details")
    col3, col4 = st.columns(2)
    
    with col3:
        proposal_date = st.date_input("Proposal Date *", value=date.today(), key="proposal_date")
        project_name = st.text_input("Project Name *", placeholder="e.g., Self Storage – 7400 22nd Ave N. St Petersburg 33710", key="project_name")
    
    with col4:
        project_address = st.text_input(
            "Project Address",
            value=st.session_state.get('lookup_address', ''),
            placeholder="e.g., 7400 22nd Ave N",
            key="project_address"
        )
        project_city_state_zip = st.text_input(
            "City, State, Zip",
            value=f"{st.session_state.get('lookup_city', '')}, FL {st.session_state.get('lookup_zip', '')}".strip(', FL '),
            placeholder="e.g., St. Petersburg, FL 33710",
            key="project_city_state_zip"
        )
    
    project_description = st.text_area(
        "Project Description / Understanding *",
        placeholder="Enter a detailed description of the project scope, objectives, and key requirements...",
        height=150,
        key="project_description"
    )
    
    st.markdown("---")
    
    # Fee Type Selection
    st.header("💰 Fee Type")
    fee_type = st.radio(
        "Select fee type for all tasks:",
        options=["Hourly, Not-to-Exceed", "Hourly", "Lump Sum"],
        horizontal=True,
        help="This fee type will apply to all selected tasks"
    )
    
    st.markdown("---")
    
    # Project Understanding Assumptions
    st.header("📋 Project Understanding Assumptions")
    st.markdown("Check the assumptions that apply to this project. These will appear in the Project Understanding section.")
    
    col_assume1, col_assume2 = st.columns(2)
    
    with col_assume1:
        assume_survey = st.checkbox(
            "Boundary, topographic, and tree survey provided by Client",
            value=True
        )
        assume_environmental = st.checkbox(
            "Environmental/Biological assessment provided"
        )
        assume_geotech = st.checkbox(
            "Geotechnical investigation report provided"
        )
        assume_zoning = st.checkbox(
            "Use is consistent with future land use and zoning",
            value=True
        )
        assume_utilities = st.checkbox(
            "Utilities available at project boundary with adequate capacity",
            value=True
        )
    
    with col_assume2:
        assume_offsite = st.checkbox(
            "Offsite roadway improvements not included",
            value=True
        )
        assume_traffic = st.checkbox(
            "Traffic study provided by others or not required",
            value=True
        )
        assume_one_phase = st.checkbox(
            "Project constructed in one (1) phase",
            value=True
        )
        
        # Conceptual plan date input
        col_plan1, col_plan2 = st.columns([1, 2])
        with col_plan1:
            has_conceptual_plan = st.checkbox("Based on conceptual plan")
        with col_plan2:
            conceptual_plan_date = st.text_input(
                "Plan Date",
                placeholder="e.g., 10/15/2024",
                disabled=not has_conceptual_plan,
                label_visibility="collapsed"
            )

# TAB 3: Scope of Services
with tab3:
    # Task Selection and Fees
    st.header("✅ Scope of Services")
st.markdown("Select the tasks to include in the proposal and enter the fee for each task.")

selected_tasks = {}

for task_num in sorted(DEFAULT_FEES.keys()):
    task = DEFAULT_FEES[task_num]
    
    # Compact row layout
    col_check, col_name, col_fee = st.columns([1, 4, 2])
    
    with col_check:
        task_selected = st.checkbox(
            f"{task_num}",
            key=f"check_{task_num}"
        )
    
    with col_name:
        st.markdown(f"**Task {task_num}: {task['name']}**")
    
    with col_fee:
        fee_amount = st.number_input(
            "Fee ($)",
            min_value=0,
            value=None,
            placeholder=f"{task['amount']:,}",
            key=f"fee_{task_num}",
            disabled=not task_selected,
            label_visibility="collapsed"
        )
    
    if task_selected:
        final_fee = fee_amount if fee_amount is not None else task['amount']
        selected_tasks[task_num] = {
            'name': task['name'],
            'fee': final_fee,
            'type': fee_type  # Use global fee type
        }
        
        # Task 310 - Construction Phase Services needs hour allocations
        if task_num == '310':
            st.markdown("**📝 Construction Phase Service Hours:**")
            col_hrs1, col_hrs2, col_hrs3 = st.columns(3)
            
            with col_hrs1:
                shop_drawing_hours = st.number_input(
                    "Shop Drawing Review (hrs)",
                    min_value=0,
                    value=30,
                    key="shop_drawing_hours"
                )
                rfi_hours = st.number_input(
                    "RFI Response (hrs)",
                    min_value=0,
                    value=50,
                    key="rfi_hours"
                )
            
            with col_hrs2:
                oac_meetings = st.number_input(
                    "OAC Meetings (count)",
                    min_value=0,
                    value=24,
                    key="oac_meetings"
                )
                site_visits = st.number_input(
                    "Site Visits (count)",
                    min_value=0,
                    value=4,
                    key="site_visits"
                )
            
            with col_hrs3:
                record_drawing_hours = st.number_input(
                    "Record Drawings (hrs)",
                    min_value=0,
                    value=40,
                    key="record_drawing_hours"
                )
                total_construction_hours = st.number_input(
                    "Total Task Hours",
                    min_value=0,
                    value=180,
                    key="total_construction_hours"
                )
            
            # Store these in the task dict
            selected_tasks[task_num]['hours'] = {
                'shop_drawing': shop_drawing_hours,
                'rfi': rfi_hours,
                'oac_meetings': oac_meetings,
                'site_visits': site_visits,
                'record_drawing': record_drawing_hours,
                'total': total_construction_hours
            }

st.markdown("---")

# TAB 4: Permitting & Billing
with tab4:
    # Permitting Requirements
    st.header("📋 Permitting Requirements")
st.markdown("Select the permits/approvals required for this project (applies to Task 150 - Civil Permitting):")

# Get permit defaults based on selected county
permit_config = PERMIT_MAPPING.get(county, {})
default_permits = permit_config.get('default_permits', [])
ahj_name = permit_config.get('ahj_name', 'Authority Having Jurisdiction')
wmd_name = permit_config.get('wmd_short', 'Water Management District')

col_permit1, col_permit2, col_permit3 = st.columns(3)

with col_permit1:
    permit_ahj = st.checkbox(
        f"{ahj_name}",
        value=("ahj" in default_permits),
        help="Primary permitting authority"
    )
    permit_sewer = st.checkbox(
        "Sewer Provider",
        value=("sewer" in default_permits)
    )
    permit_water = st.checkbox(
        "Water Provider",
        value=("water" in default_permits)
    )

with col_permit2:
    permit_wmd_erp = st.checkbox(
        f"{wmd_name} ERP",
        value=("wmd_erp" in default_permits),
        help="Environmental Resources Permit"
    )
    permit_fdep = st.checkbox(
        "FDEP Potable Water/Wastewater",
        value=("fdep" in default_permits)
    )
    permit_fdot_drainage = st.checkbox(
        "FDOT Drainage Connection",
        value=("fdot_drainage" in default_permits)
    )

with col_permit3:
    permit_fdot_driveway = st.checkbox(
        "FDOT Driveway Connection",
        value=("fdot_driveway" in default_permits)
    )
    permit_fdot_utility = st.checkbox(
        "FDOT Utility Connection",
        value=("fdot_utility" in default_permits)
    )
    permit_fema = st.checkbox(
        "FEMA",
        value=("fema" in default_permits)
    )

st.markdown("---")

# Section 6: Summary
if selected_tasks:
    st.header("📊 Selected Tasks Summary")
    
    total_fee = 0
    for task_num in sorted(selected_tasks.keys()):
        task = selected_tasks[task_num]
        st.write(f"✓ Task {task_num}: {task['name']} — **${task['fee']:,}**")
        total_fee += task['fee']
    
    st.markdown("---")
    st.markdown(f"### **Total Fee: ${total_fee:,}**")
    st.markdown("---")
else:
    st.info("👆 Select at least one task to generate a proposal")
    
    st.markdown("---")
    
    # Invoice & Billing Information (still in tab4)
    st.header("📧 Invoice & Billing Information")

col_inv1, col_inv2 = st.columns(2)

with col_inv1:
    invoice_email = st.text_input(
        "Invoice Email Address",
        placeholder="e.g., accounting@company.com",
        help="Primary email for invoices"
    )
    
    kh_signer_name = st.text_input(
        "Kimley-Horn Signer Name",
        placeholder="e.g., John Smith, PE"
    )

with col_inv2:
    invoice_cc_email = st.text_input(
        "CC Email (optional)",
        placeholder="e.g., manager@company.com",
        help="Additional recipient for invoices"
    )
    
    kh_signer_title = st.text_input(
        "Kimley-Horn Signer Title",
        placeholder="e.g., Senior Project Manager"
    )

# TAB 5: Generate Document
with tab5:
    st.header("📄 Generate Proposal")
    
    required_fields = {
        'County': st.session_state.get('county', ''),
        'City': st.session_state.get('city', ''),
        'Client Name': st.session_state.get('client_name', ''),
        'Contact Person': st.session_state.get('contact_person', ''),
        'Address Line 1': st.session_state.get('address_line1', ''),
        'Address Line 2': st.session_state.get('address_line2', ''),
        'Project Name': st.session_state.get('project_name', ''),
        'Project Description': st.session_state.get('project_description', '')
    }
    
    missing_fields = [field for field, value in required_fields.items() if not value]
    
    if missing_fields:
        st.warning(f"⚠️ Please fill in: {', '.join(missing_fields)}")
    
    if not selected_tasks:
        st.warning("⚠️ Please select at least one task")
    
    can_generate = not missing_fields and bool(selected_tasks)
    
    if st.button("🚀 Generate Proposal Document", type="primary", disabled=not can_generate):
        with st.spinner("Generating proposal document..."):
            try:
                # Collect assumptions
                assumptions = []
                if assume_survey:
                    assumptions.append("Boundary, topographic, and tree survey will be provided by the Client.")
                if assume_environmental:
                    assumptions.append("An Environmental/Biological assessment and Geotechnical investigation report will be provided by the Client.")
                if assume_geotech:
                    assumptions.append("A Geotechnical investigation report will be provided by the Client.")
                if assume_zoning:
                    assumptions.append("The proposed use is consistent with the property's future land use and zoning designations.")
                if has_conceptual_plan and conceptual_plan_date:
                    assumptions.append(f"This proposal is based on the conceptual site plan dated {conceptual_plan_date}.")
                if assume_utilities:
                    assumptions.append("Utilities are available at the project boundary and have the capacity to serve the proposed development.")
                if assume_offsite:
                    assumptions.append("Offsite roadway improvements or right-of-way permitting is not included.")
                if assume_traffic:
                    assumptions.append("Traffic Study, impact analysis, and traffic counts, if required, will be provided by others.")
                if assume_one_phase:
                    assumptions.append("The project will be constructed in one (1) phase.")
                
                # Collect permitting requirements
                permit_config = PERMIT_MAPPING.get(county, {})
                ahj_name = permit_config.get('ahj_name', 'Authority Having Jurisdiction')
                wmd_full = permit_config.get('wmd', 'Water Management District')
                
                permits = []
                if permit_ahj:
                    permits.append(ahj_name)
                if permit_sewer:
                    permits.append("Sewer Provider" if not ahj_name else f"{ahj_name} Sewer")
                if permit_water:
                    permits.append("Water Provider" if not ahj_name else f"{ahj_name} Water")
                if permit_wmd_erp:
                    permits.append(f"{wmd_full} Environmental Resources Permit (ERP)")
                if permit_fdep:
                    permits.append("Florida Department of Environmental Protection (FDEP) Potable Water and Wastewater Permit")
                if permit_fdot_drainage:
                    permits.append("FDOT Drainage Connection Permit")
                if permit_fdot_driveway:
                    permits.append("FDOT Driveway Connection Permit")
                if permit_fdot_utility:
                    permits.append("FDOT Utility Connection Permit")
                if permit_fema:
                    permits.append("FEMA")
                
                client_info = {
                    'name': st.session_state.get('client_name', ''),
                    'contact': st.session_state.get('contact_person', ''),
                    'address1': st.session_state.get('address_line1', ''),
                    'address2': st.session_state.get('address_line2', ''),
                    'phone': st.session_state.get('phone', ''),
                    'email': st.session_state.get('email', '')
                }
                
                project_info = {
                    'date': st.session_state.get('proposal_date', date.today()).strftime("%B %d, %Y"),
                    'name': st.session_state.get('project_name', ''),
                    'address': st.session_state.get('project_address', ''),
                    'city_state_zip': st.session_state.get('project_city_state_zip', ''),
                    'description': st.session_state.get('project_description', ''),
                    'county': st.session_state.get('county', ''),
                    'city': st.session_state.get('city', ''),
                    'parcel_id': st.session_state.get('parcel_id', ''),
                    'fee_type': fee_type,
                    'permits': permits
                }
                
                buffer = BytesIO()
                temp_path = '/tmp/temp_proposal.docx'
                generate_proposal_document(client_info, project_info, selected_tasks, assumptions, permits, temp_path)
                
                with open(temp_path, 'rb') as f:
                    buffer.write(f.read())
                buffer.seek(0)
                
                filename = f"Proposal_{project_name.replace(' ', '_')[:30]}_{proposal_date.strftime('%Y%m%d')}.docx"
                
                st.success("✅ **Proposal document generated successfully!**")
                
                st.download_button(
                    label="📥 Download Word Document",
                    data=buffer.getvalue(),
                    file_name=filename,
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    type="primary",
                    use_container_width=True
                )
                
            except Exception as e:
                st.error(f"❌ **Error:** {str(e)}")
                with st.expander("Show Error Details"):
                    st.exception(e)

st.markdown("---")
st.caption("Development Services Proposal Generator | Kimley-Horn")
